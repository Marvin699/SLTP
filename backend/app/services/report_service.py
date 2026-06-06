"""方案报告生成服务 - 生成Word和PDF报告"""
import json
import os
import re
import subprocess
import tempfile
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

_CJK_FONT_REGISTERED = False
_CJK_FONT_NAME = None

def _register_cjk_font():
    global _CJK_FONT_REGISTERED, _CJK_FONT_NAME
    if _CJK_FONT_REGISTERED:
        return _CJK_FONT_NAME
    candidates = [
        ('NotoSansSC', '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc'),
        ('NotoSansSC2', '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc'),
        ('WenQuanYi', '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc'),
        ('SourceHanSansSC', '/usr/share/fonts/opentype/noto/SourceHanSansSC-Regular.otf'),
        ('PingFang', '/System/Library/Fonts/PingFang.ttc'),
        ('STHeiti', '/System/Library/Fonts/STHeiti Medium.ttc'),
        ('ArialUnicode', '/Library/Fonts/Arial Unicode.ttf'),
        ('ArialUnicode2', '/System/Library/Fonts/Supplemental/Arial Unicode.ttf'),
        ('SimSun', 'C:/Windows/Fonts/simsun.ttc'),
        ('SimHei', 'C:/Windows/Fonts/simhei.ttf'),
        ('MicrosoftYaHei', 'C:/Windows/Fonts/msyh.ttc'),
    ]
    for name, path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                _CJK_FONT_REGISTERED = True
                _CJK_FONT_NAME = name
                return name
            except Exception:
                continue
    _CJK_FONT_REGISTERED = True
    _CJK_FONT_NAME = 'Helvetica'
    return 'Helvetica'


def _set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_run_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _add_heading(doc, text, level=1):
    """添加标题"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    if level == 1:
        _set_run_font(run, '黑体', 16, bold=True)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        _set_run_font(run, '黑体', 14, bold=True)
    else:
        _set_run_font(run, '黑体', 12, bold=True)
    paragraph.space_after = Pt(12)
    paragraph.space_before = Pt(12)
    return paragraph


def _add_paragraph(doc, text, bold=False, align='left', font_size=10.5):
    """添加段落"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    _set_run_font(run, '宋体', font_size, bold)
    if align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.space_after = Pt(6)
    return paragraph


def _add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        _set_cell_shading(cell, '4472C4')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                _set_run_font(run, '宋体', 10, bold=True)
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 填充数据
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _set_run_font(run, '宋体', 9)
            # 隔行变色
            if row_idx % 2 == 1:
                _set_cell_shading(cell, 'D6E4F0')

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    doc.add_paragraph()  # 表格后空一行
    return table


def generate_report_data(task: dict, solution: dict, diagnosis: dict = None) -> dict:
    """
    生成报告数据
    
    参数:
        task: 任务配置
        solution: 方案数据
        diagnosis: 诊断结果（可选）
    
    返回:
        报告数据结构
    """
    # 提取基本信息
    depot = task.get('depot', {})
    demand_points = task.get('demand_points', [])
    uavs = task.get('uavs', [])
    
    # 项目名称
    project_name = f"{depot.get('name', '未知')}无人机应急物资配送项目"
    
    # 无人机型号汇总
    uav_types = {}
    for uav in uavs:
        name = uav.get('name', '未知')
        if name not in uav_types:
            uav_types[name] = 0
        uav_types[name] += 1
    uav_summary = '+'.join([f"{name}×{count}" for name, count in uav_types.items()])
    
    # 检查是否有冷链物资
    has_cold_chain = False
    total_weight = 0
    for dp in demand_points:
        materials = dp.get('materials', [])
        for m in materials:
            weight = m.get('weight', 0)
            total_weight += weight
            material_type = m.get('type', '')
            if any(keyword in material_type for keyword in ['冷链', '冷藏', '胰岛素', '疫苗']):
                has_cold_chain = True
    
    # 副标题
    subtitle = uav_summary
    if has_cold_chain:
        subtitle += "，含冷链"
    
    # 方案类型（可自定义，默认路径规划方案）
    scheme_type = task.get('scheme_type', '路径规划方案')
    
    # 汇总统计
    summary = solution.get('summary', {})
    stats = {
        'total_distance': summary.get('total_distance', 0),
        'total_time': summary.get('total_time', 0),
        'total_trips': summary.get('total_trips', 0),
        'drone_count': summary.get('drone_count', 0),
        'village_count': summary.get('village_count', 0),
        'total_weight': round(total_weight, 2),
        'feasible': summary.get('feasible', True),
    }
    
    # 路径汇总表
    route_table = solution.get('route_table', [])
    
    # 村庄配送详情表
    village_table = solution.get('village_table', [])
    
    # 无人机配送详情表
    drone_table = solution.get('drone_table', [])
    
    # 诊断评分
    scores = {}
    if diagnosis and diagnosis.get('four_dimensional_scores'):
        scores = diagnosis.get('four_dimensional_scores', {})
    
    return {
        'project_name': project_name,
        'scheme_type': scheme_type,
        'subtitle': subtitle,
        'depot': depot,
        'demand_points': demand_points,
        'uavs': uavs,
        'stats': stats,
        'route_table': route_table,
        'village_table': village_table,
        'drone_table': drone_table,
        'scores': scores,
        'diagnosis': diagnosis,
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
    }


def generate_word_report(report_data: dict, output_path: str = None) -> str:
    """
    生成Word报告
    
    参数:
        report_data: 报告数据
        output_path: 输出路径（可选）
    
    返回:
        生成的文件路径
    """
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 封面
    _add_heading(doc, report_data['project_name'], level=1)
    _add_heading(doc, report_data['scheme_type'], level=1)
    _add_paragraph(doc, f"（{report_data['subtitle']}）", align='center', font_size=12)
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 生成时间
    _add_paragraph(doc, f"生成时间：{report_data['generated_at']}", align='center')
    doc.add_page_break()
    
    # 一、项目概况
    _add_heading(doc, '一、项目概况', level=2)
    
    depot = report_data['depot']
    stats = report_data['stats']
    
    overview_text = f"""
    本次无人机应急物资配送任务以{depot.get('name', '未知')}为配送中心，
    共需配送{stats['village_count']}个村庄，总物资重量约{stats['total_weight']}kg。
    调配{stats['drone_count']}架无人机执行配送任务，
    预计总飞行距离{stats['total_distance']}km，总飞行时间{stats['total_time']}分钟，
    共需执行{stats['total_trips']}趟次。
    """
    _add_paragraph(doc, overview_text.strip())
    
    # 配送点列表
    _add_paragraph(doc, '配送点信息：', bold=True)
    demand_points = report_data['demand_points']
    dp_rows = []
    for i, dp in enumerate(demand_points):
        materials = dp.get('materials', [])
        material_names = '、'.join([m.get('type', '物资') for m in materials]) if materials else '无'
        dp_rows.append([
            i + 1,
            dp.get('name', '未知'),
            dp.get('priority', '普通'),
            round(sum(m.get('weight', 0) for m in materials), 2),
            material_names,
        ])
    
    _add_table(doc,
        ['序号', '村庄名称', '优先级', '需求重量(kg)', '物资类型'],
        dp_rows,
        [1.5, 3, 2.5, 2.5, 6]
    )
    
    # 无人机配置
    _add_paragraph(doc, '无人机配置：', bold=True)
    uavs = report_data['uavs']
    uav_rows = []
    for i, uav in enumerate(uavs):
        uav_rows.append([
            i + 1,
            uav.get('name', '未知'),
            uav.get('max_payload', 0),
            uav.get('range_km', uav.get('max_range', 0)),
            uav.get('max_speed', 0),
        ])
    
    _add_table(doc,
        ['序号', '机型', '最大载重(kg)', '最大航程(km)', '最大速度(km/h)'],
        uav_rows,
        [1.5, 4, 2.5, 2.5, 3]
    )
    
    # 诊断评分（如果有）
    scores = report_data.get('scores', {})
    if scores:
        _add_paragraph(doc, '方案评估：', bold=True)
        score_rows = [
            ['安全性', scores.get('safety', 0), '载重安全、航程安全、飞行安全'],
            ['时效性', scores.get('timeliness', 0), '配送时间、无人机利用率、优先级满足'],
            ['经济性', scores.get('economy', 0), '路径优化、负载均衡、资源利用'],
            ['可行性', scores.get('feasibility', 0), '需求覆盖、硬约束满足、可执行性'],
        ]
        _add_table(doc,
            ['评估维度', '得分', '评估说明'],
            score_rows,
            [3, 2, 8]
        )
    
    doc.add_page_break()
    
    # 二、航线规划与调度
    _add_heading(doc, '二、航线规划与调度', level=2)
    
    # 汇总统计
    _add_paragraph(doc, '无人机应急物资配送路径汇总表', bold=True, align='center')
    summary_rows = [
        ['总飞行距离', f"{stats['total_distance']} km"],
        ['总飞行时间', f"{stats['total_time']} 分钟"],
        ['总趟次', str(stats['total_trips'])],
        ['无人机数量', str(stats['drone_count'])],
        ['配送村庄数', str(stats['village_count'])],
    ]
    _add_table(doc, ['指标', '数值'], summary_rows, [5, 5])
    
    # 路径汇总表
    _add_paragraph(doc, '路径汇总', bold=True)
    route_table = report_data.get('route_table', [])
    if route_table:
        route_rows = []
        for r in route_table:
            route_rows.append([
                r.get('route_id', '-'),
                r.get('drone_name', '-'),
                r.get('route_path', '-'),
                r.get('route_nodes', '-'),
                r.get('distance', 0),
                r.get('time', 0),
                r.get('weight', 0),
                r.get('village_name', '-'),
                r.get('trip_count', 1),
            ])
        _add_table(doc,
            ['路径编号', '无人机', '路径', '途径节点', '距离(km)', '时间(min)', '配送重量(kg)', '目标村庄', '趟次'],
            route_rows,
            [1.2, 2, 2.5, 2.5, 1.5, 1.5, 2, 2.5, 1]
        )
    
    # 村庄配送详情
    _add_paragraph(doc, '村庄配送详情', bold=True)
    village_table = report_data.get('village_table', [])
    if village_table:
        village_rows = []
        for v in village_table:
            village_rows.append([
                v.get('village_id', '-'),
                v.get('village_name', '-'),
                v.get('demand_weight', 0),
                v.get('drone_name', '-'),
                v.get('drone_weight', 0),
                v.get('trip_count', 0),
                v.get('one_way_distance', 0),
                v.get('special_req', '普通配送'),
            ])
        _add_table(doc,
            ['村庄编号', '村庄名称', '需求重量(kg)', '配送无人机', '无人机配送重量(kg)', '趟次', '单程距离(km)', '特殊要求'],
            village_rows,
            [1.2, 2, 2, 2.5, 2.5, 1.2, 2, 3]
        )
    
    # 无人机配送详情
    _add_paragraph(doc, '无人机配送详情', bold=True)
    drone_table = report_data.get('drone_table', [])
    if drone_table:
        drone_rows = []
        for d in drone_table:
            drone_rows.append([
                d.get('drone_id', '-'),
                d.get('drone_type', '-'),
                d.get('speed_ms', 0),
                d.get('max_payload', 0),
                d.get('total_distance', 0),
                d.get('total_time', 0),
                d.get('total_trips', 0),
                d.get('villages', '-'),
                d.get('note', '-'),
            ])
        _add_table(doc,
            ['无人机编号', '机型', '速度(m/s)', '最大载重(kg)', '总飞行距离(km)', '总飞行时间(min)', '总趟次', '服务村庄', '备注'],
            drone_rows,
            [1.5, 2, 1.5, 2, 2, 2, 1.2, 3, 2]
        )
    
    # 保存文件
    if not output_path:
        output_path = tempfile.mktemp(suffix='.docx')
    
    doc.save(output_path)
    return output_path


def _sanitize_filename(s: str) -> str:
    return re.sub(r'[\\/:*?"<>|\s]+', '_', s).strip('_') or 'report'


def convert_word_to_pdf(word_path: str, pdf_path: str = None) -> str:
    """
    将Word转换为PDF

    策略:
      1) 优先尝试 libreoffice (服务器需 apt install libreoffice fonts-noto-cjk)
         → Word 原样转 PDF, 中文完美, 图片/表格/页眉页脚全部保留
      2) 回退: 用 reportlab + 中文字体 基于 Word 内容渲染 PDF
         (只保证文字/表格, 不保证图片和样式; 建议装 libreoffice)
    """
    import shutil as _shutil

    if not pdf_path:
        pdf_path = word_path.replace('.docx', '.pdf')

    soffice = _shutil.which('libreoffice') or _shutil.which('soffice')
    if soffice:
        try:
            out_dir = os.path.dirname(pdf_path) or '.'
            subprocess.run(
                [soffice, '--headless', '--norestore', '--nologo', '--convert-to', 'pdf',
                 '--outdir', out_dir, word_path],
                check=True, timeout=180,
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
            converted = os.path.join(out_dir, os.path.splitext(os.path.basename(word_path))[0] + '.pdf')
            if os.path.exists(converted) and os.path.abspath(converted) != os.path.abspath(pdf_path):
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                os.rename(converted, pdf_path)
            if os.path.exists(pdf_path):
                print(f"[report] convert_word_to_pdf -> libreoffice ({soffice}) -> {pdf_path}")
                return pdf_path
        except Exception as e:
            print(f"[report] libreoffice conversion failed ({soffice}): {e}, falling back")

    font = _register_cjk_font()
    print(f"[report] libreoffice not installed, falling back to reportlab (font={font})")
    return _render_pdf_fallback(word_path, pdf_path)


def _safe_text(v) -> str:
    if v is None:
        return ''
    if isinstance(v, (dict, list)):
        try:
            return json.dumps(v, ensure_ascii=False)
        except Exception:
            return str(v)
    return str(v)


def _render_pdf_fallback(word_path: str, pdf_path: str) -> str:
    """libreoffice 不可用时, 用 reportlab 渲染包含中文正文的 PDF."""
    font = _register_cjk_font()

    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(word_path)
    except Exception:
        doc = None

    reportlab_doc = SimpleDocTemplate(pdf_path, pagesize=A4,
                                     rightMargin=2 * cm, leftMargin=2 * cm,
                                     topMargin=2 * cm, bottomMargin=2 * cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'],
                                fontName=font, fontSize=20, alignment=1, spaceAfter=16)
    h1_style = ParagraphStyle('H1', parent=styles['Heading2'],
                              fontName=font, fontSize=16, alignment=0, spaceBefore=10, spaceAfter=8)
    h2_style = ParagraphStyle('H2', parent=styles['Heading3'],
                              fontName=font, fontSize=13, alignment=0, spaceBefore=8, spaceAfter=6)
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
                                fontName=font, fontSize=11, leading=16)

    story = []
    story.append(Paragraph("智慧低空应急运输路径规划报告", title_style))
    story.append(Paragraph(f"Word源文件: {os.path.basename(word_path)}",
                           ParagraphStyle('Muted', parent=body_style, textColor=colors.grey, fontSize=9)))
    story.append(Spacer(1, 12))

    if doc is not None:
        for para in doc.paragraphs:
            txt = _safe_text(para.text).strip()
            if not txt:
                continue
            style = body_style
            lvl = para.style.name if para.style else ''
            if 'Heading 1' in lvl or lvl == 'Title':
                style = h1_style
            elif 'Heading 2' in lvl:
                style = h2_style
            story.append(Paragraph(txt.replace('\n', '<br/>'), style))

        if doc.tables:
            for ti, table in enumerate(doc.tables):
                story.append(Paragraph(f"表 {ti + 1}", h2_style))
                data = []
                for ri, row in enumerate(table.rows):
                    data.append([Paragraph(_safe_text(c.text).strip().replace('\n', ' '),
                                          ParagraphStyle('TC', parent=body_style, fontSize=10))
                                 for c in row.cells])
                if data:
                    t = Table(data, repeatRows=1)
                    t.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (-1, -1), font),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('GRID', (0, 0), (-1, -1), 0.3, colors.grey),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e8f1ff')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    story.append(Spacer(1, 4))
                    story.append(t)
                    story.append(Spacer(1, 10))
    else:
        story.append(Paragraph("（无法解析Word文件，请尝试用 libreoffice 转换）", body_style))

    reportlab_doc.build(story)
    return pdf_path


def generate_report(task: dict, solution: dict, diagnosis: dict = None, 
                   output_dir: str = None) -> Dict[str, str]:
    """
    生成完整报告（Word + PDF）
    
    参数:
        task: 任务配置
        solution: 方案数据
        diagnosis: 诊断结果（可选）
        output_dir: 输出目录（可选）
    
    返回:
        {'word': word路径, 'pdf': pdf路径, 'data': 报告数据}
    """
    # 生成报告数据
    report_data = generate_report_data(task, solution, diagnosis)
    
    # 创建输出目录
    if not output_dir:
        output_dir = tempfile.gettempdir()
    os.makedirs(output_dir, exist_ok=True)
    
    # 生成文件名
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    project_name = _sanitize_filename(report_data['project_name'].replace(' ', '_'))
    base_name = f"{project_name}_{timestamp}"
    
    # 生成Word
    word_path = os.path.join(output_dir, f"{base_name}.docx")
    generate_word_report(report_data, word_path)
    
    # 生成PDF
    pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
    convert_word_to_pdf(word_path, pdf_path)
    
    return {
        'word': word_path,
        'pdf': pdf_path,
        'data': report_data,
        'filename': base_name,
    }
