from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 设置中文字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def set_run_font(run, font_name='微软雅黑', size=11, bold=False, color=None, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_custom(text, level=1):
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_run_font(run, size=18 if level == 1 else 14, bold=True, color=RGBColor(0x1e, 0x3a, 0x8c) if level == 1 else RGBColor(0x2c, 0x52, 0x84))
    paragraph.paragraph_format.space_before = Pt(18) if level == 1 else Pt(14)
    paragraph.paragraph_format.space_after = Pt(10)
    return paragraph

def add_title(text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=24, bold=True, color=RGBColor(0x1e, 0x3a, 0x8c))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(20)
    paragraph.paragraph_format.space_after = Pt(20)
    return paragraph

def add_subtitle(text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=12, color=RGBColor(0x55, 0x55, 0x55))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(20)
    return paragraph

def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(10)
    return p

def add_formula_box(formula_text, description=None):
    """添加公式展示框"""
    # 公式框背景色（浅蓝灰）
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4F8')
    pPr.append(shd)
    
    # 左侧边框
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '3B82F6')
    pBdr.append(left)
    pPr.append(pBdr)
    
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.3
    
    run = p.add_run(formula_text)
    set_run_font(run, font_name='Cambria Math', size=15, bold=True, color=RGBColor(0x1e, 0x3a, 0x8c), italic=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if description:
        desc_p = doc.add_paragraph()
        desc_run = desc_p.add_run(description)
        set_run_font(desc_run, size=10, color=RGBColor(0x66, 0x66, 0x66), italic=True)
        desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_p.paragraph_format.space_after = Pt(8)
    
    return p

def add_tip_box(text):
    """添加提示框"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFFBEB')
    pPr.append(shd)
    
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'F59E0B')
    pBdr.append(left)
    pPr.append(pBdr)
    
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.5
    
    icon_run = p.add_run('💡 ')
    set_run_font(icon_run, size=11)
    
    tip_run = p.add_run('直观理解：')
    set_run_font(tip_run, size=11, bold=True, color=RGBColor(0x92, 0x40, 0x00))
    
    text_run = p.add_run(text)
    set_run_font(text_run, size=11, color=RGBColor(0x66, 0x66, 0x66))
    return p

def add_normal_text(text, size=11, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 表头背景色
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2C5282')
        tcPr.append(shd)
    
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            set_run_font(run, size=10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 偶数行背景
            if r_idx % 2 == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F7FAFC')
                tcPr.append(shd)
    
    return table

# ============================
# 开始生成文档内容
# ============================

add_title('蚁群算法（ACO）核心公式说明')
add_subtitle('Ant Colony Optimization — Mathematical Formulation')
add_hr()

# 一、算法概述
add_heading_custom('一、算法概述', level=1)
add_normal_text('蚁群算法（Ant Colony Optimization，简称 ACO）是一种模仿蚂蚁觅食行为的启发式优化算法，通过模拟蚂蚁在路径上释放信息素、并根据信息素浓度选择路径的机制，寻找最优解。')
add_normal_text('本系统采用蚁群算法解决多无人机路径规划问题，目标是在满足载重、航程约束的前提下，使所有无人机的总配送距离最短。')

add_hr()

# 二、核心公式
add_heading_custom('二、核心公式', level=1)

# 公式1
add_heading_custom('公式 1：状态转移概率（选择下一个需求点）', level=2)
add_normal_text('蚂蚁在当前节点 i 选择下一个节点 j 的概率计算如下：')

add_formula_box(
    'P(j) = τᵢⱼᵅ × ηᵢⱼᵝ',
    '（公式 1-1）状态转移概率'
)

add_normal_text('参数说明：', size=11, bold=True, color=RGBColor(0x2c, 0x52, 0x84))
add_table(
    ['符号', '含义', '说明', '默认值'],
    [
        ['τᵢⱼ (tau_ij)', '信息素浓度', '边 (i,j) 上积累的"经验值"', '动态变化'],
        ['α (alpha)', '信息素重要度', '控制历史经验的权重', '1.0'],
        ['ηᵢⱼ (eta_ij)', '启发式函数值', '基于距离、优先级等计算的"吸引力"', '多因子计算'],
        ['β (beta)', '启发函数重要度', '控制贪心倾向的权重', '5.0'],
    ]
)

add_tip_box('蚂蚁倾向选择信息素浓度高（其他蚂蚁走过）且启发值大（更近、更重要）的路径。')

add_hr()

# 公式2
add_heading_custom('公式 2：启发式函数（多因子吸引力）', level=2)
add_normal_text('启发式函数综合考虑距离、优先级、载重、航程四个因素，计算候选点的吸引力：')

add_formula_box(
    'η = 0.4 × η_dist + 0.3 × η_priority + 0.15 × η_load + 0.15 × η_range',
    '（公式 2-1）综合启发式函数'
)

add_normal_text('各因子计算方式：', size=11, bold=True, color=RGBColor(0x2c, 0x52, 0x84))
add_table(
    ['因子名称', '计算公式', '权重', '含义'],
    [
        ['距离因子 η_dist', '1 / distance', '0.4', '距离越近，吸引力越大'],
        ['优先级因子 η_priority', 'priority_weight / 5.0', '0.3', '优先级越高，吸引力越大'],
        ['载重因子 η_load', '1 - load_ratio', '0.15', '载重越轻，飞行越灵活'],
        ['航程因子 η_range', 'remaining_range / max_range', '0.15', '剩余航程越充裕，安全性越高'],
    ]
)

add_normal_text('优先级权重映射表：', size=11, bold=True, color=RGBColor(0x2c, 0x52, 0x84))
add_table(
    ['优先级等级', '权重值'],
    [
        ['紧急（urgent）', '5.0'],
        ['高（high）', '3.0'],
        ['中（medium）', '1.0'],
        ['低（low）', '0.5'],
        ['普通（normal）', '0.3'],
    ]
)

add_tip_box('蚂蚁优先选择距离近、优先级高、自身载重轻、剩余航程充裕的需求点。')

add_hr()

# 公式3
add_heading_custom('公式 3：信息素挥发', level=2)
add_normal_text('每轮迭代结束后，所有路径上的信息素按一定比例衰减，避免算法陷入局部最优：')

add_formula_box(
    'τᵢⱼ = (1 - ρ) × τᵢⱼ',
    '（公式 3-1）信息素挥发'
)

add_normal_text('参数说明：', size=11, bold=True, color=RGBColor(0x2c, 0x52, 0x84))
add_table(
    ['符号', '含义', '默认值', '取值范围建议'],
    [
        ['ρ (rho)', '挥发率', '0.5', '0.3 ~ 0.7'],
    ]
)

add_tip_box('ρ 值越大，旧信息素消散越快，算法收敛越快，但容易陷入局部最优；ρ 值越小，算法探索能力越强，但收敛较慢。')

add_hr()

# 公式4
add_heading_custom('公式 4：信息素沉积（普通蚂蚁）', level=2)
add_normal_text('每只蚂蚁完成一次完整路径后，在经过的路径上沉积信息素：')

add_formula_box(
    'Δτᵢⱼ = Q / cost',
    '（公式 4-1）单条路径信息素沉积量'
)

add_formula_box(
    'τᵢⱼ = τᵢⱼ + Δτᵢⱼ',
    '（公式 4-2）信息素更新'
)

add_normal_text('参数说明：', size=11, bold=True, color=RGBColor(0x2c, 0x52, 0x84))
add_table(
    ['符号', '含义', '默认值'],
    [
        ['Q', '信息素强度常数', '100'],
        ['cost', '路径成本（总距离）', '动态计算'],
    ]
)

add_tip_box('路径越短（成本越低），沉积的信息素越多，鼓励其他蚂蚁选择更优的短路径。')

add_hr()

# 公式5
add_heading_custom('公式 5：精英蚂蚁强化', level=2)
add_normal_text('对本轮表现最优的若干只"精英蚂蚁"进行额外强化，放大其信息素沉积量，加快算法收敛：')

add_formula_box(
    'Δτ_elite = (Q / cost) × elite_bonus',
    '（公式 5-1）精英蚂蚁信息素强化'
)

add_normal_text('参数说明：', size=11, bold=True, color=RGBColor(0x2c, 0x52, 0x84))
add_table(
    ['符号', '含义', '默认值'],
    [
        ['elite_bonus', '精英强化倍数', '2.0（本轮最优蚂蚁）\n1.5（其他精英蚂蚁）'],
    ]
)

add_tip_box('表现最好的蚂蚁留下更明显的"标记"，引导其他蚂蚁快速向最优解收敛。')

add_hr()

# 三、算法流程
add_heading_custom('三、算法流程', level=1)

flow_text = """
开始
  ↓
初始化：信息素矩阵、参数设置
  ↓
┌─────────────────────────────────────┐
│         对每只蚂蚁循环                │
│   ┌─────────────────────────────┐   │
│   │ 1. 选择下一个节点            │   │
│   │    公式：P = τ^α × η^β       │   │
│   │    方法：轮盘赌选择           │   │
│   └─────────────────────────────┘   │
│   ┌─────────────────────────────┐   │
│   │ 2. 检查约束条件              │   │
│   │    • 载重限制：can_carry()   │   │
│   │    • 航程限制：can_reach()   │   │
│   └─────────────────────────────┘   │
│   ┌─────────────────────────────┐   │
│   │ 3. 生成完整路径              │   │
│   │    循环直到所有需求点已访问    │   │
│   └─────────────────────────────┘   │
│   ┌─────────────────────────────┐   │
│   │ 4. 计算路径成本（总距离）    │   │
│   └─────────────────────────────┘   │
└─────────────────────────────────────┘
  ↓
信息素挥发：τ = (1 - ρ) × τ      （公式 3-1）
  ↓
普通蚂蚁沉积：Δτ = Q / cost      （公式 4-1）
  ↓
精英蚂蚁强化：Δτ = (Q / cost) × bonus  （公式 5-1）
  ↓
是否达到最大迭代次数？
  ├── 是 → 输出全局最优解
  └── 否 → 回到"对每只蚂蚁循环"
"""

p = doc.add_paragraph()
run = p.add_run(flow_text)
set_run_font(run, font_name='Consolas', size=10, color=RGBColor(0x37, 0x41, 0x51))

# 设置流程文本背景
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'F1F5F9')
pPr.append(shd)

pBdr = OxmlElement('w:pBdr')
for side in ['top', 'bottom', 'left', 'right']:
    elem = OxmlElement(f'w:{side}')
    elem.set(qn('w:val'), 'single')
    elem.set(qn('w:sz'), '4')
    elem.set(qn('w:space'), '1')
    elem.set(qn('w:color'), 'CBD5E1')
    pBdr.append(elem)
pPr.append(pBdr)

p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)

add_hr()

# 四、参数配置说明
add_heading_custom('四、参数配置说明', level=1)
add_table(
    ['参数名称', '符号', '默认值', '调优建议'],
    [
        ['蚂蚁数量', 'num_ants', '30', '需求点 > 15 时增至 50，> 25 时增至 80'],
        ['最大迭代次数', 'max_iterations', '100', '需求点 > 15 时增至 150，> 25 时增至 200'],
        ['信息素重要度', 'α (alpha)', '1.0', '值越大越依赖历史经验，建议范围 0.5~2.0'],
        ['启发函数重要度', 'β (beta)', '5.0', '值越大越贪心，建议范围 2.0~10.0'],
        ['挥发率', 'ρ (rho)', '0.5', '值越大收敛越快但易陷入局部最优，建议 0.3~0.7'],
        ['信息素强度', 'Q', '100', '影响信息素沉积量，建议范围 50~200'],
        ['精英蚂蚁数', 'elite_ants', '3', '额外强化的精英数量，建议 2~5'],
    ]
)

add_hr()

# 五、约束条件
add_heading_custom('五、求解约束条件', level=1)

constraints = [
    ('载重约束', '无人机装载的物资重量不得超过最大载重：current_load + point_weight ≤ max_payload'),
    ('航程约束', '无人机必须能在剩余航程内到达下一个需求点并返回基地：distance + return_distance ≤ remaining_range'),
    ('能耗约束', '无人机的电池能量足以完成当前任务，能耗 = f(distance, load)'),
]

for idx, (name, desc) in enumerate(constraints, 1):
    p = doc.add_paragraph()
    
    # 编号图标
    num_run = p.add_run(f'  {idx}.  ')
    set_run_font(num_run, size=12, bold=True, color=RGBColor(0x2c, 0x52, 0x84))
    
    # 约束名
    name_run = p.add_run(f'{name}：')
    set_run_font(name_run, size=11, bold=True, color=RGBColor(0x1a, 0x20, 0x2c))
    
    # 描述
    desc_run = p.add_run(desc)
    set_run_font(desc_run, size=11, color=RGBColor(0x4b, 0x55, 0x63))
    
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

add_hr()

# 文档尾
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1 = footer_p.add_run('文档版本：V1.0    ')
set_run_font(run1, size=10, color=RGBColor(0x99, 0x99, 0x99))
run2 = footer_p.add_run('生成日期：2026年7月')
set_run_font(run2, size=10, color=RGBColor(0x99, 0x99, 0x99))

doc.save('/Users/wy/Desktop/智慧低空应急运输教学平台/蚁群算法公式说明.docx')
print('Word文档已生成：蚁群算法公式说明.docx')
