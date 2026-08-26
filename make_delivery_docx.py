from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_center_title(text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(15)
    paragraph.paragraph_format.space_after = Pt(15)
    return paragraph

def add_info_row(label, value, bold=True):
    paragraph = doc.add_paragraph()
    run1 = paragraph.add_run(label)
    run1.font.name = '微软雅黑'
    run1.font.size = Pt(11)
    run1.font.bold = bold
    run2 = paragraph.add_run(value)
    run2.font.name = '微软雅黑'
    run2.font.size = Pt(11)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph

def add_section_title(text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(12)
    run.font.bold = True
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    return paragraph

def add_normal_text(text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(11)
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.5
    return paragraph

def add_bullet(text):
    paragraph = doc.add_paragraph(text, style='List Bullet')
    for run in paragraph.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(11)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.5
    return paragraph

def add_seal_area(title):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.font.name = '微软雅黑'
    run.font.size = Pt(12)
    run.font.bold = True
    paragraph.paragraph_format.space_before = Pt(15)
    paragraph.paragraph_format.space_after = Pt(8)
    
    seal = doc.add_paragraph()
    seal_run = seal.add_run('▇▇▇▇▇▇▇▇▇▇')
    seal_run.font.name = '微软雅黑'
    seal_run.font.size = Pt(24)
    seal_run.font.color.rgb = RGBColor(139, 0, 0)
    seal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    note = doc.add_paragraph()
    note_run = note.add_run('（此处加盖红色公章）')
    note_run.font.name = '微软雅黑'
    note_run.font.size = Pt(10)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph

add_center_title('供货单')

add_info_row('单据编号：', 'ZYZY-2026-0708')
add_info_row('资金项目：', '现代职业教育质量提升计划专项资金（课程资源建设类）')
add_info_row('项目名称：', '低空应急物流课程图谱智能体平台搭建技术支持服务')
add_info_row('供货单位：', 'XX 数字教育科技有限公司')
add_info_row('统一社会信用代码：', '91XXXXXXXXXXXXXXXXXX')
add_info_row('联系电话：', '13XXXXXXXXX')
add_info_row('收货单位：', 'XX 职业技术学院 交通物流学院')
add_info_row('收货地址：', 'XX 市 XX 区 XX 大道 XX 号')
add_info_row('供货完成日期：', '2026 年 07 月 21 日')

add_section_title('')

table = doc.add_table(rows=4, cols=6)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['序号', '产品/服务名称（财务合规品名）', '单位', '数量', '单价(元)', '小计(元)']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = '微软雅黑'
            run.font.size = Pt(10)
            run.font.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.2
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

items = [
    ['1', '低空应急物流职业岗位能力标准适配技术服务', '项', '1', '2000', '2000'],
    ['2', '三层课程图谱数字化交互功能开发技术服务', '项', '1', '3000', '3000'],
    ['3', '平台部署与系统调试优化技术服务', '项', '1', '1500', '1500'],
]

for row_idx, item in enumerate(items):
    for col_idx, text in enumerate(item):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(10)
            if col_idx in [0, 2, 3, 4, 5]:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing = 1.2
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

add_section_title('功能参数')
add_bullet('1、依托企业智能体平台技术，适配学院自主调研梳理的低空应急物流职业岗位能力标准、课程知识点、实操技能点及安全规范要点，提供定制化平台技术适配支撑；')
add_bullet('2、基于学院自主设计的知识图谱、能力图谱、问题图谱三层可视化课程图谱框架，提供平台部署、数字化交互功能开发、系统调试优化等技术服务，完成图谱可视化展示、智能交互、资源联动等数字化功能落地，保障自主研发课程图谱正常教学应用。')

add_section_title('金额合计')
add_normal_text('合计小写：￥6500.00')
add_normal_text('合计大写：人民币陆仟伍佰元整')

add_section_title('备注')
add_normal_text('本单据所有供货内容均为职业教育数字化课程教学资源技术服务，归属现代职业教育质量提升计划专项资金课程资源建设范畴，不含硬件设备、商用智能系统开发；')
add_normal_text('课程图谱核心内容、体系架构、知识技能体系均由学院自主开发，供货方仅提供平台技术适配与开发支持服务；')
add_normal_text('所有资源仅限本校低空应急物流专业课堂、实训教学使用，不对外商业化运营。')

add_seal_area('【供货单位盖章区】')
add_info_row('经办人签字：', '__________    联系电话：__________', bold=False)
add_info_row('日期：', '2026 年 07 月 21 日', bold=False)

add_seal_area('【收货验收单位盖章区】')
add_info_row('验收负责人签字：', '__________    专业负责人签字：__________', bold=False)
add_info_row('验收日期：', '2026 年 07 月 21 日', bold=False)

add_seal_area('【财务审核栏】')
add_info_row('财务审核意见：', '□符合专项资金列支要求，准予入账', bold=False)
add_info_row('审核人签字：', '__________    审核日期：__________', bold=False)

doc.save('/Users/wy/Desktop/智慧低空应急运输教学平台/项目交付清单.docx')
print('Word文档已生成：项目交付清单.docx')
